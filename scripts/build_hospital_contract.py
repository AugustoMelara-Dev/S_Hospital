from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DIR = Path("docs/contratos")
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT_FILE = OUT_DIR / "Contrato_S_Hospital_borrador.docx"


BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
GRAY_FILL = "F2F4F7"
LIGHT_BLUE_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
BORDER = "C8D1DC"


def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=BORDER):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True)
        r = p.add_run(text[len(bold_prefix):])
        set_font(r)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.add_run(text)
    set_font(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    r = p.add_run(text)
    set_font(r)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        set_font(run, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True,
                 color=BLUE if level < 3 else DARK_BLUE)
    return p


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [6.5])
    set_table_borders(table)
    cell = table.cell(0, 0)
    shade_cell(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_font(r, bold=True, color=DARK_BLUE)
    p.add_run("\n")
    r = p.add_run(body)
    set_font(r)
    doc.add_paragraph()


def add_label_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=2)
    set_table_width(table, [1.85, 4.65])
    set_table_borders(table)
    for i, (label, value) in enumerate(rows):
        c0, c1 = table.rows[i].cells
        shade_cell(c0, GRAY_FILL)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r = p0.add_run(label)
        set_font(r, bold=True, color=DARK_BLUE)
        p1 = c1.paragraphs[0]
        r = p1.add_run(value)
        set_font(r)
    doc.add_paragraph()
    return table


def add_matrix(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    set_table_borders(table)
    for idx, h in enumerate(headers):
        cell = table.rows[0].cells[idx]
        shade_cell(cell, LIGHT_BLUE_FILL)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_font(r, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            p = cells[idx].paragraphs[0]
            r = p.add_run(value)
            set_font(r, size=10 if len(value) > 80 else 11)
            set_cell_margins(cells[idx])
    doc.add_paragraph()
    return table


def configure_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167


def build_contract():
    doc = Document()
    configure_styles(doc)

    footer = doc.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("Borrador contractual S_Hospital - revisar con abogado antes de firmar")
    set_font(r, size=9, color=RGBColor(89, 89, 89))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CONTRATO DE IMPLEMENTACION, LICENCIA DE USO Y MANTENIMIENTO")
    set_font(r, size=16, bold=True, color=DARK_BLUE)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Sistema Hospitalario Local S_Hospital")
    set_font(r, size=13, bold=True)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Borrador editable - Honduras - version 1.0")
    set_font(r, size=10, color=RGBColor(89, 89, 89))

    add_callout(
        doc,
        "Uso recomendado",
        "Este borrador esta pensado para que la titular comercial frente al hospital sea tu jefa o su empresa. "
        "Augusto Jose Melara Amaya puede figurar como desarrollador autorizado, contacto tecnico o subcontratista. "
        "Si el hospital contrata directamente contigo, cambia la Proveedora por tu nombre y elimina las clausulas de reventa.",
    )

    add_heading(doc, "1. Partes", 1)
    add_label_table(doc, [
        ("Proveedora", "[Nombre completo de tu jefa o razon social], con identidad/RTN [____], domicilio en [____]."),
        ("Cliente", "[Nombre legal del hospital], con RTN [____], domicilio en [____]."),
        ("Desarrollador autorizado", "Augusto Jose Melara Amaya, quien podra ejecutar soporte tecnico, desarrollo, instalacion y mantenimiento por cuenta de la Proveedora, segun autorizacion interna."),
        ("Fecha de vigencia", "[____ de __________ de 2026]."),
    ])

    add_para(doc, "La Proveedora y el Cliente podran denominarse conjuntamente como las Partes. El Desarrollador autorizado no sera parte obligada frente al Cliente salvo que firme expresamente como proveedor directo, fiador, garante o responsable contractual.")

    add_heading(doc, "2. Antecedentes y objeto", 1)
    add_para(doc, "El Cliente requiere un sistema local para caja hospitalaria, facturacion, pagos, reportes, catalogo de servicios, respaldos y recibos institucionales. La Proveedora ofrece implementar y licenciar el uso del sistema S_Hospital, junto con servicios de instalacion, configuracion, capacitacion basica, garantia limitada y mantenimiento opcional.")
    add_para(doc, "El objeto de este contrato es regular la entrega, uso, soporte y mantenimiento del sistema S_Hospital en una red local LAN, sin dependencia de internet para la operacion ordinaria de login, facturacion, reportes o impresion.")

    add_heading(doc, "3. Descripcion del sistema", 1)
    add_bullet(doc, "Aplicacion hospitalaria local con frontend React + TypeScript, API Laravel y base de datos MySQL/MariaDB.")
    add_bullet(doc, "Operacion en servidor local dentro de la red LAN del Cliente; los equipos cliente acceden por navegador mediante IP local o dominio interno.")
    add_bullet(doc, "Modulos previstos: usuarios y permisos, catalogo de servicios, caja, facturacion, pagos, anulaciones, reportes, backups y recibos institucionales PDF/papel.")
    add_bullet(doc, "El sistema no reemplaza expediente clinico completo ni sistemas medicos especializados; el paciente requiere como dato minimo el nombre para factura.")

    add_heading(doc, "4. Alcance incluido", 1)
    add_matrix(doc, ["Area", "Incluido"], [
        ("Instalacion", "Despliegue en una computadora servidor local, configuracion inicial de base de datos, backend, frontend y servidor web segun ambiente disponible."),
        ("Configuracion", "Carga inicial de catalogo de servicios, usuarios base, roles iniciales, metodos de pago y datos institucionales de recibos."),
        ("Facturacion y caja", "Crear factura, registrar pagos, asociar factura pagada a caja, cajero, metodo de pago y fecha, reimprimir recibos y consultar reportes."),
        ("Recibos", "Recibo institucional principal en carta, media carta o A5. Formatos 80mm/58mm solo como compatibilidad secundaria si se acuerda."),
        ("Backups", "Configuracion de backup diario automatico y backup manual desde panel admin cuando el ambiente tecnico lo permita."),
        ("Capacitacion", "Una capacitacion basica de hasta [__] horas para cajeros/admins sobre uso ordinario del sistema."),
    ], [1.45, 5.05])

    add_heading(doc, "5. Reglas de negocio no negociables", 1)
    for item in [
        "Las facturas deben emitirse con nombre del paciente.",
        "No se borran facturas; las anulaciones requieren permiso, motivo y auditoria.",
        "Toda factura pagada debe quedar asociada a caja, cajero, metodo de pago y fecha.",
        "El recibo principal no debe exponer QR, codigos de barras ni codigos internos, salvo acuerdo escrito distinto.",
        "La eritropoyetina se maneja como medicamento de L.25 y sera gratuita cuando se marque paciente con receta de dialisis.",
        "Las facturas historicas no se recalculan desde catalogos actuales; se conservan snapshots de precio y nombre en los items de factura.",
        "El backend sera la fuente de verdad de totales, impuestos, pagos, anulaciones y numeracion; el frontend solo podra previsualizar.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "6. Precio, forma de pago y entrega", 1)
    add_label_table(doc, [
        ("Precio de implementacion", "L. 12,000.00, salvo que las Partes pacten otro monto por escrito."),
        ("Forma de pago sugerida", "50% al iniciar y 50% contra entrega funcional; editable segun acuerdo comercial."),
        ("Entrega funcional", "Se considera entregado cuando el Cliente puede crear factura, cobrar, imprimir/reimprimir recibo y consultar reporte basico desde el navegador en la LAN."),
        ("Cambios posteriores", "Nuevas funciones, reportes especiales, integraciones, migraciones complejas o cambios de alcance se cotizaran por separado mediante orden de cambio."),
    ])

    add_heading(doc, "7. Garantia limitada", 1)
    add_para(doc, "Durante treinta (30) dias calendario contados desde la entrega funcional, la Proveedora corregira sin costo adicional errores atribuibles al software entregado que impidan el flujo principal contratado. La garantia no cubre cambios de criterio, nuevas funciones, fallas de hardware, virus, perdida por mal uso, problemas electricos, danos al servidor, manipulacion directa de la base de datos, redes mal configuradas o software de terceros no administrado por la Proveedora.")

    add_heading(doc, "8. Mantenimiento opcional", 1)
    add_para(doc, "Finalizada la garantia, el mantenimiento sera opcional y no se entendera incluido indefinidamente en el precio de implementacion. Si el Cliente desea soporte continuo, debera contratar una modalidad mensual o soporte por evento.")
    add_matrix(doc, ["Plan", "Cuota sugerida", "Incluye"], [
        ("Basico", "L. 1,500/mes", "Soporte remoto o telefonico en horario laboral, correccion de errores menores, revision de backups y hasta [__] horas mensuales."),
        ("Operativo", "L. 2,500/mes", "Todo lo basico, prioridad en caja/facturacion, pequenos ajustes de recibo/reportes y hasta [__] horas mensuales."),
        ("Critico", "L. 4,000/mes", "Prioridad alta, visita mensual o asistencia presencial pactada, revision preventiva y hasta [__] horas mensuales."),
        ("Por evento", "Desde L. 500 por caso", "Atencion puntual sin mensualidad; visitas, urgencias y cambios se cotizan segun distancia, horario y complejidad."),
    ], [1.35, 1.35, 3.80])
    add_para(doc, "Los saldos de horas no usadas no se acumulan salvo pacto escrito. Nuevos modulos, integraciones, cambios fiscales, reportes a medida, migraciones de datos y redisenos se facturan aparte.")

    add_heading(doc, "9. Niveles de soporte", 1)
    add_matrix(doc, ["Prioridad", "Ejemplo", "Respuesta objetivo"], [
        ("Alta", "No se puede facturar, cobrar o abrir caja.", "Mismo dia laboral o dentro de [4-8] horas habiles si existe plan mensual activo."),
        ("Media", "Reporte incorrecto, recibo con ajuste menor, usuario bloqueado.", "Dentro de [1-2] dias habiles."),
        ("Baja", "Mejora visual, consulta, configuracion menor no urgente.", "Dentro de [3-5] dias habiles o en la siguiente ventana de mantenimiento."),
    ], [1.2, 3.4, 1.9])

    add_heading(doc, "10. Obligaciones del Cliente", 1)
    for item in [
        "Proveer servidor, energia, red local, navegadores compatibles e impresoras en buen estado.",
        "Designar usuarios responsables y custodiar contrasenas.",
        "No compartir accesos administrativos con personal no autorizado.",
        "Permitir respaldos y no apagar abruptamente el servidor durante procesos criticos.",
        "No modificar directamente archivos del sistema o base de datos sin autorizacion tecnica.",
        "Reportar incidentes con fecha, usuario, descripcion, captura o numero de factura cuando aplique.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "11. Propiedad intelectual y licencia", 1)
    add_para(doc, "Salvo pacto escrito distinto, S_Hospital se licencia para uso interno del Cliente y no se vende el codigo fuente ni la propiedad intelectual. El Cliente recibe derecho de uso no exclusivo, no transferible y limitado a sus instalaciones o red local autorizada.")
    add_para(doc, "El Cliente no podra revender, sublicenciar, publicar, copiar para terceros, descompilar, modificar el codigo fuente ni entregar el sistema a otro proveedor sin autorizacion escrita. Los datos generados por la operacion del hospital pertenecen al Cliente.")
    add_para(doc, "La Proveedora declara que cuenta con autorizacion suficiente del Desarrollador autorizado para comercializar, instalar y mantener el sistema bajo este contrato. Cualquier disputa interna entre Proveedora y Desarrollador no debera afectar la continuidad razonable del servicio contratado por el Cliente.")

    add_heading(doc, "12. Confidencialidad y datos", 1)
    add_para(doc, "Las Partes mantendran confidencial toda informacion tecnica, operativa, comercial, de pacientes, usuarios, reportes, credenciales, respaldos y datos de facturacion a los que tengan acceso. La informacion solo se usara para ejecutar este contrato.")
    add_para(doc, "El sistema no esta disenado para almacenar expediente clinico completo. Si el Cliente decide ingresar datos adicionales no previstos, sera responsable de sus politicas internas, permisos y cumplimiento normativo aplicable.")

    add_heading(doc, "13. Seguridad, backups y continuidad", 1)
    add_bullet(doc, "La Proveedora configurara mecanismos razonables de backup local segun el ambiente disponible.")
    add_bullet(doc, "El Cliente debe conservar respaldos externos o copias fuera del servidor si desea proteccion ante robo, incendio, dano fisico o perdida total del equipo.")
    add_bullet(doc, "La Proveedora no garantiza continuidad si el servidor, red, sistema operativo, disco duro, electricidad o impresoras fallan por causas ajenas al software.")

    add_heading(doc, "14. Limitacion de responsabilidad", 1)
    add_para(doc, "La responsabilidad total de la Proveedora por reclamos relacionados con el sistema, soporte o mantenimiento no excedera el monto efectivamente pagado por el Cliente durante los tres (3) meses anteriores al evento que origino el reclamo, salvo dolo comprobado o prohibicion legal aplicable.")
    add_para(doc, "La Proveedora no sera responsable por lucro cesante, perdida indirecta, dano reputacional, sanciones causadas por uso indebido, errores de datos ingresados por usuarios, falta de backups externos, fallas de hardware, red, energia, internet, impresoras o software de terceros.")

    add_heading(doc, "15. Reventa y autorizacion comercial", 1)
    add_para(doc, "Cuando la Proveedora comercialice el sistema desarrollado o mantenido por Augusto Jose Melara Amaya, debera conservar un acuerdo interno vigente que le permita ofrecerlo al Cliente. El Cliente podra tratar a la Proveedora como responsable comercial principal, salvo que Augusto firme un contrato directo con el Cliente.")
    add_para(doc, "Si la Proveedora deja de operar, incumple pagos internos o pierde autorizacion del Desarrollador, las Partes podran acordar una transicion tecnica ordenada para no afectar la operacion del hospital, respetando saldos pendientes y derechos de propiedad intelectual.")

    add_heading(doc, "16. Terminacion", 1)
    add_para(doc, "Cualquiera de las Partes podra terminar el mantenimiento mensual con aviso escrito de treinta (30) dias. La terminacion del mantenimiento no elimina el derecho de uso del sistema ya pagado, salvo incumplimiento grave de licencia, reventa no autorizada, falta de pago o uso indebido.")
    add_para(doc, "A la terminacion, el Cliente debera pagar saldos vencidos. La Proveedora podra entregar una copia final de respaldo de la base de datos en formato tecnico razonable, siempre que no existan restricciones legales, de seguridad o saldos pendientes que lo impidan.")

    add_heading(doc, "17. Ley aplicable y solucion de conflictos", 1)
    add_para(doc, "Este contrato se interpretara conforme a las leyes aplicables de la Republica de Honduras, salvo que las Partes acuerden por escrito otra jurisdiccion valida. Antes de acudir a acciones formales, las Partes procuraran resolver diferencias mediante reunion, acta de acuerdos y plazo razonable de subsanacion.")

    add_heading(doc, "18. Aceptacion y firmas", 1)
    add_para(doc, "Leido el presente contrato, las Partes manifiestan su aceptacion en la fecha indicada.")
    add_matrix(doc, ["Parte", "Nombre, cargo, firma y fecha"], [
        ("Proveedora", "Nombre: ______________________________\nCargo: _______________________________\nFirma: _______________________________\nFecha: _______________________________"),
        ("Cliente", "Nombre: ______________________________\nCargo: _______________________________\nFirma: _______________________________\nFecha: _______________________________"),
        ("Aceptacion tecnica opcional", "Augusto Jose Melara Amaya\nRol: Desarrollador autorizado / soporte tecnico\nFirma: _______________________________\nFecha: _______________________________"),
    ], [1.8, 4.7])

    add_heading(doc, "Anexo A - Criterio de aceptacion funcional", 1)
    for item in [
        "Crear una factura con nombre de paciente y al menos un servicio.",
        "Registrar pago con caja, cajero, metodo de pago y fecha.",
        "Imprimir o generar recibo institucional principal.",
        "Reimprimir recibo de factura pagada.",
        "Consultar reporte basico de caja o facturacion por fecha.",
        "Probar backup manual y confirmar ubicacion de backup automatico si aplica.",
    ]:
        add_number(doc, item)

    add_heading(doc, "Anexo B - Exclusiones", 1)
    for item in [
        "Expediente clinico completo, farmacia avanzada, inventario medico completo o integraciones contables no descritas.",
        "Compra, reparacion o garantia de computadoras, servidores, impresoras, red, UPS, routers o licencias de terceros.",
        "Soporte fuera de horario, visitas urgentes, viaticos, migracion masiva de datos historicos o recuperacion forense.",
        "Cambios exigidos por ley, autoridad fiscal, administracion hospitalaria o nuevas politicas despues de la aceptacion funcional.",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Anexo C - Acuerdo interno recomendado entre Proveedora y Desarrollador", 1)
    add_callout(
        doc,
        "No entregar este anexo al hospital si contiene condiciones privadas.",
        "Este anexo protege la relacion entre tu jefa y Augusto. Puede separarse en otro documento antes de compartir el contrato principal con el Cliente.",
    )
    add_label_table(doc, [
        ("Titular comercial", "[Nombre de tu jefa o empresa], quien firma y cobra al hospital."),
        ("Desarrollador", "Augusto Jose Melara Amaya."),
        ("Pago por implementacion", "La Proveedora pagara al Desarrollador L. [____] por la implementacion o el porcentaje [____]% de lo cobrado al Cliente."),
        ("Mantenimiento", "Si el hospital paga mantenimiento, la Proveedora pagara al Desarrollador L. [____] mensual o [____]% de la mensualidad por soporte tecnico."),
        ("Autorizacion de reventa", "El Desarrollador autoriza a la Proveedora a revender el sistema solo al Cliente indicado y bajo los limites de licencia del contrato principal."),
        ("No competencia desleal", "La Proveedora no podra entregar codigo, instaladores o backups a terceros para copiar, revender o modificar el sistema sin autorizacion escrita del Desarrollador."),
        ("Continuidad", "Si la Proveedora deja de pagar al Desarrollador, este podra suspender nuevos trabajos despues de aviso escrito, sin afectar datos propiedad del Cliente."),
    ])

    add_heading(doc, "Anexo D - Referencias consultadas para estructura", 1)
    add_para(doc, "Estas referencias fueron usadas solo como guia de estructura contractual comun: alcance de mantenimiento, licencia, soporte, exclusiones, confidencialidad, limitacion de responsabilidad y terminacion. No sustituyen revision legal local.")
    for item in [
        "PandaDoc - Software Maintenance Agreement Template: https://www.pandadoc.com/software-maintenance-agreement-template/",
        "Business in a Box - Software Maintenance Agreement overview: https://www.business-in-a-box.com/template/software-maintenance-agreement-D805/",
        "SEC archive - Software License and Maintenance Agreement example: https://www.sec.gov/Archives/edgar/data/855109/000085510904000007/exh1036.htm",
        "AI Lawyer - Source Code License Agreement Template overview: https://ailawyer.pro/templates/source-code-license-agreement-template",
        "Koley Jessen - Software license agreement/IP considerations: https://www.koleyjessen.com/insights/publications/form-software-license-agreement-intellectual-property",
    ]:
        add_bullet(doc, item)

    doc.save(OUT_FILE)
    return OUT_FILE


if __name__ == "__main__":
    print(build_contract().resolve())
